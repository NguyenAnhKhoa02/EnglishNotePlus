import os
import docx
import subprocess

DOCX_DIR = "\\".join([os.path.dirname(os.path.dirname(__file__)),"docx"])

def getPathFileDocx(nameFile : str) -> str:
    return os.path.abspath("\\".join([DOCX_DIR,nameFile.lower()]))

def deleteParagraphs(nameFile : str):
    doc = docx.Document(open(getPathFileDocx(nameFile),'rb'))
    all_paragraphs = doc.paragraphs
    
    for para in all_paragraphs:
        p = para._element
        p.getparent().remove(p)
        p._p = p._element = None
    doc.save(open(getPathFileDocx(nameFile),'wb'))    

def openEmptyMicrosoftWord(namefile: str) -> None:
    deleteParagraphs(namefile)

    os.startfile(getPathFileDocx(namefile))

def openMicrosoftWordWithText(nameFile : str, text : str) -> None:
    deleteParagraphs(nameFile)
    
    doc = docx.Document(open(getPathFileDocx(nameFile),"rb"))
    doc.add_paragraph(text)
    doc.save(open(getPathFileDocx(nameFile),'wb'))
    
    os.startfile(getPathFileDocx(nameFile))
    
def closeMicrosoftWord():
    pid = int(str(subprocess.getoutput("wmic process get processid,name|find \"WINWORD\"")).split("WINWORD.EXE")[1].strip())
    
    subprocess.call(f"TASKKILL /F /PID {pid}", shell=True)
    
def getTextFromDocx(nameFile : str)-> str:
    f = open(getPathFileDocx(nameFile),'rb')
    strFromDocx = ""
    for para in  docx.Document(f).paragraphs:
        strFromDocx += para.text + "\n"
        
    return strFromDocx.strip()